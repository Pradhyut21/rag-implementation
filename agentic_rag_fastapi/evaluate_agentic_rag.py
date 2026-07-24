import json
import os
import time

from docx import Document
import requests

# ─────────────────────────────────────────────────────────────
# 15-Query Evaluation Harness
# Covers: factual recall, comparative, procedural, edge-case,
#         missing-info detection, multi-hop, and open-ended queries
# ─────────────────────────────────────────────────────────────

BASE_URL = os.getenv("EVAL_BASE_URL", "http://127.0.0.1:8002")
API_KEY = os.getenv("API_KEY", "demo-rag-2026")
HEADERS = {"X-API-Key": API_KEY}


def create_eval_document() -> str:
    filename = "eval_rag_doc.docx"
    doc = Document()
    doc.add_heading("Agentic RAG Architecture — Comprehensive Overview", level=0)

    sections = [
        (
            "1. Introduction",
            "Agentic Retrieval-Augmented Generation (Agentic RAG) is a design pattern that introduces "
            "multiple autonomous agents to orchestrate the retrieval, reasoning, and synthesis phases in RAG systems. "
            "Unlike vanilla RAG which simply retrieves documents and outputs answers in a single step, "
            "Agentic RAG utilises a planner, query rewriter, sufficient context checker, and feedback loop "
            "to ensure high-fidelity, grounded answers. The system achieves this without hallucinating "
            "by explicitly verifying whether retrieved context is sufficient before synthesis.",
        ),
        (
            "2. Core Agents",
            "The Planner Agent is the entry point, decomposing complex queries into 2-4 targeted sub-questions. "
            "The Query Rewriter Agent transforms each sub-question into a dense-retrieval-optimised search query. "
            "The Sufficient Context Agent (SC Agent) evaluates context sufficiency using three evidence types: "
            "explicit (exact answer present), partial (topic covered but detail missing), and missing (not in context). "
            "Finally, the Synthesis Agent drafts the grounded final answer using only the retrieved evidence.",
        ),
        (
            "3. The Feedback Loop",
            "When the SC Agent flags insufficient context (partial or missing evidence), it generates a "
            "structured feedback_log and missing_information list. These are passed through the Query Rewriter "
            "to produce targeted retrieval queries for the next iteration. The system supports up to 2 iterations "
            "before falling back to best-available synthesis. This self-correcting loop improves context coverage "
            "by an average of 34% compared to single-pass retrieval.",
        ),
        (
            "4. Reasoning Modes",
            "Standard Mode executes the 5-phase pipeline: Plan → Rewrite → Fanout → SC-Check → Synthesise. "
            "Chain of Thought (CoT) Mode runs 6 sequential reasoning stages with explicit intermediate outputs. "
            "Tree of Thought (ToT) Mode generates 3 parallel reasoning branches, scores each across 5 dimensions "
            "(retrieval similarity, coverage, completeness, evidence quality, confidence), selects the best branch, "
            "and merges the top 2 branches if their scores are within 0.05 of each other.",
        ),
        (
            "5. Observability",
            "The system uses monkey-patching to intercept all LLM calls and FAISS operations transparently. "
            "Every request is assigned a session_id via Python ContextVar propagation. "
            "Telemetry is stored in a 10-table SQLite schema including sessions, spans, events, errors, "
            "tokens, latency, CoT stages, ToT branches, branch scores, and branch evaluations. "
            "Average end-to-end latency for standard mode is 8.3 seconds. "
            "CoT mode adds approximately 3.2 seconds overhead. "
            "ToT mode adds approximately 6.7 seconds overhead versus standard mode. "
            "The observability dashboard exposes sessions, traces, events, errors, token costs, and latency breakdown.",
        ),
        (
            "6. Vector Storage",
            "Documents are chunked into 6-sentence windows with 2-sentence overlap using NLTK sentence tokenisation. "
            "Each chunk is embedded using sentence-transformers all-MiniLM-L6-v2 (384 dimensions). "
            "FAISS IndexFlatIP with cosine normalisation is used for similarity search. "
            "Indexes are persisted as .index and .pkl files. "
            "In-memory caching of loaded VectorStore objects reduces repeated disk I/O. "
            "The system supports PDF and DOCX uploads up to 20MB with automatic OCR fallback for scanned PDFs.",
        ),
        (
            "7. Security",
            "The API is protected by an X-API-Key header (configurable via API_KEY environment variable). "
            "Rate limiting is enforced: 30 requests per minute for query endpoints, 20 per hour for uploads. "
            "CORS is restricted to configured allowed origins (not wildcard). "
            "File uploads are validated by extension and PDF magic bytes. "
            "Filenames are sanitised to prevent directory traversal attacks. "
            "Registry writes use atomic temp-file replacement to prevent corruption. "
            "A thread lock protects the in-memory vector store cache from race conditions.",
        ),
        (
            "8. Deployment",
            "The system is containerised using Docker with a multi-stage build. "
            "A docker-compose.yml orchestrates the FastAPI backend (port 8002) and Vite frontend (port 5173). "
            "Environment configuration uses a .env file with GROQ_API_KEY, API_KEY, and ALLOWED_ORIGINS. "
            "The backend uses uvicorn with 2 workers. "
            "Health check endpoint at /health returns status, version, and feature list.",
        ),
    ]

    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)

    doc.save(filename)
    print(f"Created evaluation document: '{filename}'")
    return filename


def upload_doc(doc_path: str) -> str:
    print("\n[1/6] Uploading evaluation document...")
    with open(doc_path, "rb") as f:
        files = {
            "file": (
                doc_path,
                f,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        r = requests.post(f"{BASE_URL}/upload-doc", files=files, headers=HEADERS, timeout=60)
    if r.status_code != 200:
        raise RuntimeError(f"Upload failed: {r.status_code} — {r.text}")
    doc_id = r.json()["doc_id"]
    print(f"    ✓ Indexed. doc_id = {doc_id}")
    return doc_id


# 15 evaluation queries covering all categories
EVAL_QUERIES = [
    # ── Factual recall ────────────────────────────────────────
    {
        "id": "F1",
        "query": "What is Agentic RAG and how does it differ from vanilla RAG?",
        "expected_sufficient": True,
        "category": "factual",
    },
    {
        "id": "F2",
        "query": "What are the three evidence types used by the Sufficient Context Agent?",
        "expected_sufficient": True,
        "category": "factual",
    },
    {
        "id": "F3",
        "query": "What embedding model is used and what is its vector dimensionality?",
        "expected_sufficient": True,
        "category": "factual",
    },
    {
        "id": "F4",
        "query": "What is the maximum file size supported for document upload?",
        "expected_sufficient": True,
        "category": "factual",
    },
    {
        "id": "F5",
        "query": "How many tables are in the observability SQLite schema?",
        "expected_sufficient": True,
        "category": "factual",
    },
    # ── Comparative ───────────────────────────────────────────
    {
        "id": "C1",
        "query": "Compare the latency overhead of CoT mode versus ToT mode.",
        "expected_sufficient": True,
        "category": "comparative",
    },
    {
        "id": "C2",
        "query": "How does the feedback loop in Agentic RAG improve over single-pass retrieval?",
        "expected_sufficient": True,
        "category": "comparative",
    },
    {
        "id": "C3",
        "query": "What is the difference between explicit, partial, and missing evidence types?",
        "expected_sufficient": True,
        "category": "comparative",
    },
    # ── Procedural ────────────────────────────────────────────
    {
        "id": "P1",
        "query": "Describe the 5 phases of the Standard reasoning mode pipeline.",
        "expected_sufficient": True,
        "category": "procedural",
    },
    {
        "id": "P2",
        "query": "How are FAISS indexes persisted and loaded across server restarts?",
        "expected_sufficient": True,
        "category": "procedural",
    },
    {
        "id": "P3",
        "query": "How does Tree of Thought mode select the winning reasoning branch?",
        "expected_sufficient": True,
        "category": "procedural",
    },
    # ── Multi-hop ─────────────────────────────────────────────
    {
        "id": "M1",
        "query": "Which agent is responsible for both query decomposition and feedback query generation, and why?",
        "expected_sufficient": True,
        "category": "multi-hop",
    },
    {
        "id": "M2",
        "query": "How does the security model interact with rate limiting and CORS configuration?",
        "expected_sufficient": True,
        "category": "multi-hop",
    },
    # ── Missing-info detection ────────────────────────────────
    {
        "id": "X1",
        "query": "What is the exact GPU memory consumption of the embedding model during inference?",
        "expected_sufficient": False,
        "category": "missing-info",
    },
    {
        "id": "X2",
        "query": "What is the Groq LPU benchmark score on MLPerf 2024?",
        "expected_sufficient": False,
        "category": "missing-info",
    },
]


def run_query(doc_id: str, q: dict, mode: str) -> dict:
    payload = {
        "query": q["query"],
        "doc_id": doc_id,
        "top_k": 5,
        "include_trace": True,
        "response_mode": "detailed",
        "reasoning_mode": mode,
    }
    try:
        r = requests.post(f"{BASE_URL}/ask-debug", json=payload, headers=HEADERS, timeout=120)
        if r.status_code == 200:
            resp = r.json()
            return {
                "id": q["id"],
                "mode": mode,
                "query": q["query"],
                "category": q["category"],
                "sufficient": resp["context_sufficient"],
                "expected_sufficient": q["expected_sufficient"],
                "correct": resp["context_sufficient"] == q["expected_sufficient"],
                "iterations": resp.get("iterations", 1),
                "fallback_used": resp.get("fallback_used", False),
                "session_id": resp.get("session_id"),
                "evidence_type": resp.get("evidence_type", "unknown"),
                "error": None,
            }
        else:
            return {
                "id": q["id"],
                "mode": mode,
                "query": q["query"],
                "category": q["category"],
                "error": f"HTTP {r.status_code}: {r.text[:200]}",
                "correct": False,
            }
    except Exception as e:
        return {
            "id": q["id"],
            "mode": mode,
            "query": q["query"],
            "category": q["category"],
            "error": str(e),
            "correct": False,
        }


def run_evaluation():
    doc_path = create_eval_document()

    try:
        doc_id = upload_doc(doc_path)
    except RuntimeError as e:
        print(f"❌ {e}")
        return

    all_results = []
    modes = ["standard", "cot", "tot"]

    for mode in modes:
        print(
            f"\n[{'2' if mode == 'standard' else '3' if mode == 'cot' else '4'}/6] Running {mode.upper()} mode ({len(EVAL_QUERIES)} queries)..."
        )
        mode_correct = 0
        for q in EVAL_QUERIES:
            result = run_query(doc_id, q, mode)
            all_results.append(result)
            status = "✓" if result.get("correct") else "✗"
            err_info = f" [ERROR: {result['error'][:60]}]" if result.get("error") else ""
            print(
                f"    [{status}] {result['id']} ({result['category']}) — sufficient={result.get('sufficient')} (expected={q['expected_sufficient']}){err_info}"
            )
            if result.get("correct"):
                mode_correct += 1
        print(
            f"    ── {mode.upper()} accuracy: {mode_correct}/{len(EVAL_QUERIES)} ({mode_correct / len(EVAL_QUERIES) * 100:.1f}%)"
        )

    # ── Metrics calculation ───────────────────────────────────
    print("\n[5/6] Calculating metrics...")

    def calc_metrics(results, mode):
        mode_res = [r for r in results if r.get("mode") == mode and not r.get("error")]
        if not mode_res:
            return {}
        correct = sum(1 for r in mode_res if r.get("correct"))
        total = len(mode_res)
        avg_iters = sum(r.get("iterations", 1) for r in mode_res) / total if total else 0
        by_cat = {}
        for cat in ["factual", "comparative", "procedural", "multi-hop", "missing-info"]:
            cat_r = [r for r in mode_res if r.get("category") == cat]
            if cat_r:
                by_cat[cat] = sum(1 for r in cat_r if r.get("correct")) / len(cat_r)
        return {
            "mode": mode,
            "total_queries": total,
            "correct": correct,
            "accuracy": round(correct / total, 3) if total else 0,
            "avg_iterations": round(avg_iters, 2),
            "by_category": by_cat,
        }

    metrics = {mode: calc_metrics(all_results, mode) for mode in modes}
    errors = [r for r in all_results if r.get("error")]

    print("\n══════════════════════════════════════════════")
    print("  EVALUATION RESULTS SUMMARY (15 queries x 3 modes)")
    print("══════════════════════════════════════════════")
    for mode in modes:
        m = metrics[mode]
        if m:
            print(
                f"  {mode.upper():8s} → accuracy: {m['accuracy'] * 100:.1f}%  avg_iterations: {m['avg_iterations']}"
            )
            for cat, acc in m.get("by_category", {}).items():
                print(f"             {cat:16s}: {acc * 100:.0f}%")
    if errors:
        print(f"\n  ⚠️  {len(errors)} query/mode combinations errored.")
    print("══════════════════════════════════════════════")

    report = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "doc_id": doc_id,
        "num_queries": len(EVAL_QUERIES),
        "modes_tested": modes,
        "metrics": metrics,
        "results": all_results,
        "errors": errors,
    }
    report_file = "evaluation_report.json"
    with open(report_file, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, default=str)
    print(f"\n[6/6] Report saved → {report_file}")

    # Cleanup
    print("\nCleaning up...")
    try:
        r = requests.delete(f"{BASE_URL}/documents/{doc_id}", headers=HEADERS, timeout=30)
        print(f"    DELETE: {r.json().get('message', r.text)}")
    except Exception as e:
        print(f"    Cleanup failed: {e}")
    if os.path.exists(doc_path):
        os.remove(doc_path)
        print(f"    Removed {doc_path}")

    print("\n✅ Evaluation complete.")


if __name__ == "__main__":
    run_evaluation()
