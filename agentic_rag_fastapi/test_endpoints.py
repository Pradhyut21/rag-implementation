import time
import requests
from docx import Document

# create a dummy document
doc = Document()
doc.add_paragraph("Google's Agentic RAG architecture involves multiple agents working together to provide grounded answers.")
doc.add_paragraph("The role of the Sufficient Context Agent is to determine if the retrieved context is sufficient, and if not, suggest feedback.")
doc.save("test_doc.docx")

print("Document created.")

# wait for uvicorn to be ready
time.sleep(5)

print("Testing /health")
try:
    r = requests.get("http://127.0.0.1:8000/health")
    print("/health status:", r.status_code)
    print("/health response:", r.json())
except Exception as e:
    print("/health failed:", e)

print("\nTesting /upload-doc")
try:
    with open("test_doc.docx", "rb") as f:
        r = requests.post("http://127.0.0.1:8000/upload-doc", files={"file": ("test_doc.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    print("/upload-doc status:", r.status_code)
    print("/upload-doc response:", r.json())
except Exception as e:
    print("/upload-doc failed:", e)

print("\nTesting /vanilla-ask")
try:
    r = requests.post("http://127.0.0.1:8000/vanilla-ask", json={"query": "Explain Google's Agentic RAG architecture and the role of the Sufficient Context Agent.", "top_k": 3})
    print("/vanilla-ask status:", r.status_code)
    print("/vanilla-ask response:", r.json())
except Exception as e:
    print("/vanilla-ask failed:", e)
