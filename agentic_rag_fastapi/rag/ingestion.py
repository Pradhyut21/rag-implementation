import logging

from docx import Document
from nltk.tokenize import sent_tokenize
from pypdf import PdfReader

logger = logging.getLogger("agentic_rag.ingestion")


def load_docx(file_path: str) -> str:
    doc = Document(file_path)
    text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text.strip())
    return "\n".join(text)


def load_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text.strip())
    return "\n".join(pages)


def load_pdf_with_ocr(file_path: str) -> str:
    """
    OCR-based PDF extraction using Unstructured.io + Tesseract fallback.
    Handles scanned PDFs and inconsistent formatting.
    """
    # Method 1: Try Unstructured.io (best quality, handles tables/layout)
    try:
        from unstructured.partition.pdf import partition_pdf

        elements = partition_pdf(
            filename=file_path,
            strategy="hi_res",
            infer_table_structure=True,
        )
        text_parts = [str(el) for el in elements if str(el).strip()]
        if text_parts:
            logger.info(
                f"Unstructured.io OCR extracted {len(text_parts)} elements from {file_path}"
            )
            return "\n".join(text_parts)
    except ImportError:
        logger.warning("unstructured not installed, falling back to pytesseract.")
    except Exception as e:
        logger.warning(f"Unstructured OCR failed: {e}. Falling back to pytesseract.")

    # Method 2: Tesseract via pdf2image + pytesseract
    try:
        import io

        from PIL import Image
        import pytesseract

        try:
            from pdf2image import convert_from_path

            images = convert_from_path(file_path, dpi=300)
        except ImportError:
            # Fallback: just run on first page using pypdf's image extraction
            reader = PdfReader(file_path)
            images = []
            for page in reader.pages:
                for img_obj in page.images:
                    images.append(Image.open(io.BytesIO(img_obj.data)))
            if not images:
                raise ValueError("No images found in PDF for OCR.")

        pages_text = []
        for img in images:
            text = pytesseract.image_to_string(img, lang="eng")
            if text.strip():
                pages_text.append(text.strip())

        if pages_text:
            logger.info(f"pytesseract OCR extracted {len(pages_text)} pages from {file_path}")
            return "\n".join(pages_text)

    except ImportError as e:
        logger.warning(f"pytesseract/PIL not available: {e}")
    except Exception as e:
        logger.warning(f"Tesseract OCR failed: {e}")

    # Method 3: Last resort — standard pypdf extraction
    logger.info("Falling back to standard pypdf text extraction.")
    return load_pdf(file_path)


def load_document(file_path: str) -> str:
    file_lower = file_path.lower()
    if file_lower.endswith(".docx"):
        return load_docx(file_path)
    elif file_lower.endswith(".pdf"):
        text = load_pdf(file_path)
        # If standard extraction yields very little text, auto-try OCR
        if len(text.strip()) < 200:
            logger.info(f"Short text ({len(text)} chars) from PDF — attempting OCR auto-fallback.")
            ocr_text = load_pdf_with_ocr(file_path)
            if len(ocr_text.strip()) > len(text.strip()):
                return ocr_text
        return text
    else:
        raise ValueError("Unsupported file type. Only .docx and .pdf are supported.")


def chunk_text(text: str, chunk_size: int = 6, overlap: int = 2) -> list:
    """
    Sentence-based chunking with configurable size and overlap.
    Automatically handles documents with very few sentences.
    """
    sentences = sent_tokenize(text)

    if not sentences:
        return []

    # For very short documents, reduce chunk_size gracefully
    effective_chunk_size = min(chunk_size, max(1, len(sentences)))
    effective_overlap = min(overlap, effective_chunk_size - 1)

    chunks = []
    start = 0
    while start < len(sentences):
        end = min(start + effective_chunk_size, len(sentences))
        chunk = " ".join(sentences[start:end]).strip()
        if chunk:
            chunks.append(chunk)
        if end == len(sentences):
            break
        start += effective_chunk_size - effective_overlap

    return chunks
