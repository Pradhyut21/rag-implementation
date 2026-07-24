"""
Shared pytest fixtures for the Agentic RAG Platform test suite.

All expensive objects (FastAPI TestClient, embedding model, vector store)
are created once per test session and reused. External calls (Groq LLM,
FAISS disk I/O) are mocked by default so unit tests run offline.
"""

from __future__ import annotations

from collections.abc import Generator
import io
import json
import os
from pathlib import Path
from unittest.mock import MagicMock, patch

from fastapi.testclient import TestClient
import numpy as np
import pytest

# ---------------------------------------------------------------------------
# Minimal environment so app imports succeed without a real .env file
# ---------------------------------------------------------------------------
os.environ.setdefault("GROQ_API_KEY", "test-groq-key-not-real")
os.environ.setdefault("API_KEY", "test-api-key-2026")
os.environ.setdefault("DEMO_MODE", "true")
os.environ.setdefault("ALLOWED_ORIGINS", "http://localhost:5173")


# ---------------------------------------------------------------------------
# FastAPI TestClient (session-scoped — created once per test run)
# ---------------------------------------------------------------------------
@pytest.fixture()
def client() -> Generator[TestClient, None, None]:
    """
    Return a FastAPI TestClient with real routing but mocked external I/O.

    The Groq client is patched at session scope so no real API calls are made.
    """
    mock_groq = MagicMock()
    mock_choice = MagicMock()
    mock_choice.message.content = '["What is RAG?", "How does the feedback loop work?"]'
    mock_groq.chat.completions.create.return_value = MagicMock(choices=[mock_choice])

    mock_st = MagicMock()
    mock_st.encode.side_effect = lambda texts, **kwargs: np.random.rand(len(texts), 384).astype(np.float32)

    with patch("agents.llm.Groq", return_value=mock_groq), patch("rag.embeddings.SentenceTransformer", return_value=mock_st):
        from app import app

        tc = TestClient(app, raise_server_exceptions=False)
        yield tc


@pytest.fixture(scope="session")
def auth_headers() -> dict[str, str]:
    """Return auth headers using the test API key."""
    return {"X-API-Key": "test-api-key-2026"}


# ---------------------------------------------------------------------------
# Mock LLM helpers
# ---------------------------------------------------------------------------
@pytest.fixture()
def mock_llm_response():
    """
    Context-manager fixture: patch ``safe_generate`` to return a
    configurable string without making a real API call.
    """

    def _make_mock(return_value: str = "mocked response"):
        return patch("agents.llm.safe_generate", return_value=return_value)

    return _make_mock


@pytest.fixture()
def mock_planner():
    """Patch planner_agent to return two deterministic sub-queries."""
    with patch(
        "agents.planner.planner_agent",
        return_value=["What is RAG?", "How does the feedback loop work?"],
    ) as m:
        yield m


@pytest.fixture()
def mock_rewriter():
    """Patch query_rewriter to return the input unchanged."""
    with patch(
        "agents.rewriter.query_rewriter",
        side_effect=lambda q: f"rewritten: {q}",
    ) as m:
        yield m


@pytest.fixture()
def mock_synthesis():
    """Patch synthesis_agent to return a predictable answer."""
    with patch(
        "agents.synthesis.synthesis_agent",
        return_value="This is a mocked synthesis answer for testing.",
    ) as m:
        yield m


@pytest.fixture()
def mock_sc_agent_sufficient():
    """Patch sufficient_context_agent to always return sufficient context."""
    sc_result = {
        "is_context_sufficient": True,
        "missing_information": [],
        "feedback_log": "",
        "reasoning_summary": "Sufficient context found in test.",
        "evidence_type": "explicit",
    }
    with patch("agents.sufficient_context.sufficient_context_agent", return_value=sc_result) as m:
        yield m


@pytest.fixture()
def mock_sc_agent_insufficient():
    """Patch sufficient_context_agent to always return insufficient context."""
    sc_result = {
        "is_context_sufficient": False,
        "missing_information": ["Specific latency metric not found"],
        "feedback_log": "Search for latency benchmarks",
        "reasoning_summary": "Partial evidence only — retry needed.",
        "evidence_type": "partial",
    }
    with patch("agents.sufficient_context.sufficient_context_agent", return_value=sc_result) as m:
        yield m


# ---------------------------------------------------------------------------
# Mock FAISS / VectorStore
# ---------------------------------------------------------------------------
@pytest.fixture()
def mock_vector_store():
    """Return a MagicMock VectorStore that yields fake retrieval results."""
    vs = MagicMock()
    vs.search.return_value = [
        {"chunk": "RAG stands for Retrieval-Augmented Generation.", "score": 0.95, "index": 0},
        {"chunk": "The feedback loop improves context coverage.", "score": 0.87, "index": 1},
        {"chunk": "FAISS is used for vector similarity search.", "score": 0.82, "index": 2},
    ]
    return vs


@pytest.fixture()
def mock_embedding_model():
    """Return a MagicMock embedding model that produces random vectors."""
    em = MagicMock()
    em.embed_texts.side_effect = lambda texts: np.random.rand(len(texts), 384).astype(np.float32)
    em.embed_query.return_value = np.random.rand(384).astype(np.float32)
    return em


# ---------------------------------------------------------------------------
# Temporary directory fixture
# ---------------------------------------------------------------------------
@pytest.fixture()
def tmp_data_dir(tmp_path: Path) -> Path:
    """
    Create a temporary data directory structure matching production layout.
    Patches the app-level directory constants so uploads go to tmp_path.
    """
    (tmp_path / "uploads").mkdir()
    (tmp_path / "indexes").mkdir()
    (tmp_path / "debug_runs").mkdir()
    return tmp_path


# ---------------------------------------------------------------------------
# Sample document fixtures
# ---------------------------------------------------------------------------
@pytest.fixture()
def sample_docx_bytes() -> bytes:
    """
    Return the bytes of a minimal valid .docx file for upload testing.
    Uses python-docx to create an in-memory document.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Test Document", level=0)
    doc.add_paragraph(
        "Agentic RAG uses multiple agents to orchestrate retrieval and synthesis. "
        "The Sufficient Context Agent evaluates whether retrieved evidence is sufficient. "
        "The Planner decomposes complex queries into focused sub-questions. "
        "FAISS is used for fast approximate nearest-neighbour search. "
        "The feedback loop runs up to two iterations before final synthesis."
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture()
def sample_pdf_bytes() -> bytes:
    """
    Return minimal PDF bytes (valid PDF magic bytes + basic structure).
    Not a real parseable PDF — used only for magic-byte validation tests.
    """
    return b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\nendobj\n%%EOF"


@pytest.fixture()
def fake_registry(tmp_path: Path) -> dict:
    """Write a fake registry.json to tmp_path and return the registry dict."""
    registry = {
        "abc12345": {
            "doc_id": "abc12345",
            "file_name": "test_document.docx",
            "uploaded_at": "2026-07-24T00:00:00Z",
            "num_chunks": 5,
            "upload_path": str(tmp_path / "uploads" / "abc12345_test_document.docx"),
            "index_path": str(tmp_path / "indexes" / "abc12345.index"),
            "chunks_path": str(tmp_path / "indexes" / "abc12345_chunks.pkl"),
            "chunk_size": 6,
            "overlap": 2,
            "embedding_model": "all-MiniLM-L6-v2",
        }
    }
    registry_path = tmp_path / "indexes" / "registry.json"
    registry_path.parent.mkdir(parents=True, exist_ok=True)
    registry_path.write_text(json.dumps(registry, indent=2))
    return registry
