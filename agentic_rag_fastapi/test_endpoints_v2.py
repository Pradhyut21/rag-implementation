import os
import time

from docx import Document
import requests


# 1. Create a dummy test document
def create_dummy_doc():
    filename = "test_doc_v2.docx"
    doc = Document()
    doc.add_paragraph(
        "Google's Agentic RAG architecture is a state-of-the-art framework that coordinates multiple agents."
    )
    doc.add_paragraph(
        "The Sufficient Context Agent checks the intermediate draft against the retrieved context to decide if information is missing."
    )
    doc.add_paragraph(
        "If context is insufficient, it logs the feedback and triggers another iteration of retrieval with rewritten search queries."
    )
    doc.add_paragraph(
        "The Synthesis Agent is responsible for composing the final answer grounded strictly on verified context."
    )
    doc.save(filename)
    print(f"Dummy document '{filename}' created.")
    return filename


def run_tests():
    filename = create_dummy_doc()
    base_url = "http://127.0.0.1:8000"

    # Wait for Uvicorn server to start if running concurrently
    print("Waiting 3 seconds for FastAPI server to start...")
    time.sleep(3)

    # A. Test /health
    print("\n--- TEST A: Health check ---")
    try:
        r = requests.get(f"{base_url}/health")
        print("GET /health status:", r.status_code)
        print("Response:", r.json())
        assert r.status_code == 200
        assert r.json() == {"status": "ok"}
    except Exception as e:
        print("Health check failed:", e)
        return

    # B. Test /upload-doc
    print("\n--- TEST B: Upload Document ---")
    doc_id = None
    try:
        with open(filename, "rb") as f:
            files = {
                "file": (
                    filename,
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            r = requests.post(f"{base_url}/upload-doc", files=files)
        print("POST /upload-doc status:", r.status_code)
        resp = r.json()
        print("Response:", resp)
        assert r.status_code == 200
        assert "doc_id" in resp
        assert resp["file_name"] == filename
        doc_id = resp["doc_id"]
    except Exception as e:
        print("Upload failed:", e)
        return

    # C. Test /documents list
    print("\n--- TEST C: List Documents ---")
    try:
        r = requests.get(f"{base_url}/documents")
        print("GET /documents status:", r.status_code)
        docs = r.json()
        print("Response:", docs)
        assert r.status_code == 200
        assert len(docs) >= 1
        assert any(d["doc_id"] == doc_id for d in docs)
    except Exception as e:
        print("Listing documents failed:", e)

    # D. Test /documents/{doc_id} metadata
    print("\n--- TEST D: Get Document Details ---")
    try:
        r = requests.get(f"{base_url}/documents/{doc_id}")
        print(f"GET /documents/{doc_id} status:", r.status_code)
        meta = r.json()
        print("Response:", meta)
        assert r.status_code == 200
        assert meta["doc_id"] == doc_id
        assert meta["file_name"] == filename
        assert "uploaded_at" in meta
        assert "num_chunks" in meta
        assert meta["chunk_size"] == 6
        assert meta["overlap"] == 2
        assert meta["embedding_model"] == "all-MiniLM-L6-v2"
    except Exception as e:
        print("Retrieving document details failed:", e)

    # E. Test /vanilla-ask
    print("\n--- TEST E: Vanilla RAG Query ---")
    try:
        payload = {
            "query": "What is the role of the Sufficient Context Agent?",
            "doc_id": doc_id,
            "top_k": 3,
        }
        r = requests.post(f"{base_url}/vanilla-ask", json=payload)
        print("POST /vanilla-ask status:", r.status_code)
        resp = r.json()
        print("Response answer preview:", resp.get("answer", "")[:150])
        assert r.status_code == 200
        assert "answer" in resp
        assert "retrieved_chunks" in resp
        assert len(resp["retrieved_chunks"]) > 0
    except Exception as e:
        print("Vanilla ask failed:", e)

    # F. Test /ask with include_trace=False
    print("\n--- TEST F: Agentic RAG without Trace ---")
    try:
        payload = {
            "query": "What is the role of the Sufficient Context Agent?",
            "doc_id": doc_id,
            "top_k": 3,
            "include_trace": False,
        }
        r = requests.post(f"{base_url}/ask", json=payload)
        print("POST /ask (include_trace=False) status:", r.status_code)
        resp = r.json()
        print("Response:")
        print("  - iterations:", resp.get("iterations"))
        print("  - context_sufficient:", resp.get("context_sufficient"))
        print("  - missing_information:", resp.get("missing_information"))
        print("  - trace:", resp.get("trace"))
        print("  - final_context:", resp.get("final_context"))

        assert r.status_code == 200
        assert resp["trace"] is None
        assert resp["final_context"] is None
    except Exception as e:
        print("Agentic RAG without trace failed:", e)

    # G. Test /ask with include_trace=True
    print("\n--- TEST G: Agentic RAG with Trace ---")
    try:
        payload = {
            "query": "What is the role of the Sufficient Context Agent?",
            "doc_id": doc_id,
            "top_k": 3,
            "include_trace": True,
        }
        r = requests.post(f"{base_url}/ask", json=payload)
        print("POST /ask (include_trace=True) status:", r.status_code)
        resp = r.json()
        print("Response keys:", list(resp.keys()))
        assert r.status_code == 200
        assert resp["trace"] is not None
        assert resp["final_context"] is not None
        assert len(resp["trace"]) > 0
    except Exception as e:
        print("Agentic RAG with trace failed:", e)

    # H. Test /ask with missing information (failure-case validation)
    print("\n--- TEST H: Agentic RAG Missing Information Query ---")
    try:
        payload = {
            "query": "What latency measurements did Google report for the Sufficient Context Agent?",
            "doc_id": doc_id,
            "top_k": 3,
            "include_trace": True,
        }
        r = requests.post(f"{base_url}/ask", json=payload)
        print("POST /ask (missing latency) status:", r.status_code)
        resp = r.json()
        print("Response details:")
        print("  - iterations:", resp.get("iterations"))
        print("  - context_sufficient:", resp.get("context_sufficient"))
        print("  - missing_information:", resp.get("missing_information"))

        assert r.status_code == 200
        assert "iterations" in resp
        assert resp["iterations"] >= 1
        assert "context_sufficient" in resp
        # Since dummy document doesn't mention latency, context should be flagged insufficient
        # and missing_information should list missing items
        if not resp["context_sufficient"]:
            assert len(resp["missing_information"]) > 0
            print(
                "Successfully verified feedback loop: Context marked insufficient and missing info populated!"
            )
        else:
            print(
                "Note: LLM judged context sufficient, which is possible but unexpected given doc text."
            )
    except Exception as e:
        print("Missing information query test failed:", e)

    # I. Test non-existent doc_id (404 check)
    print("\n--- TEST I: Query Invalid Document ---")
    try:
        payload = {"query": "Hello", "doc_id": "nonexistent_id", "top_k": 3}
        r = requests.post(f"{base_url}/ask", json=payload)
        print("POST /ask (invalid doc_id) status:", r.status_code)
        print("Response:", r.json())
        assert r.status_code == 404
        assert "detail" in r.json()
    except Exception as e:
        print("Invalid document test failed:", e)

    # J. Test /delete-document/{doc_id}
    print("\n--- TEST J: Delete Document ---")
    try:
        r = requests.delete(f"{base_url}/documents/{doc_id}")
        print(f"DELETE /documents/{doc_id} status:", r.status_code)
        print("Response:", r.json())
        assert r.status_code == 200

        # Verify 404 on subsequent get
        r_get = requests.get(f"{base_url}/documents/{doc_id}")
        print("GET deleted doc status:", r_get.status_code)
        assert r_get.status_code == 404
        print("Delete verification complete.")
    except Exception as e:
        print("Delete document test failed:", e)

    # Cleanup local test doc
    if os.path.exists(filename):
        os.remove(filename)
        print(f"Cleaned up local file: {filename}")


if __name__ == "__main__":
    run_tests()
